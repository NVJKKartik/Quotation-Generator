import os
import re
import io
from pathlib import Path
import docx
import fitz  # PyMuPDF
import markdown
from PIL import Image
from docx.document import Document as DocxDocumentObject # Renaming to avoid conflict with docx.Document()
from docx.table import Table as DocxTableObject
from docx.text.paragraph import Paragraph as DocxParagraphObject
# Required for iterating through document body elements
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx2pdf import convert as docx_to_pdf_converter # For DOCX to PDF


def convert_pdf_to_md(pdf_path):
    """
    Convert a PDF file to Markdown format.
    
    Args:
        pdf_path (str): Path to the PDF file
        
    Returns:
        str: Markdown content
    """
    md_content = []
    
    try:
        # Open the PDF file
        doc = fitz.open(pdf_path)
        
        # Process each page
        for page_num, page in enumerate(doc):
            # Extract text
            text = page.get_text()
            
            # Add page header
            md_content.append(f"## Page {page_num + 1}\n")
            
            # Add text content
            md_content.append(text)
            md_content.append("\n\n")
            
            # Extract images
            image_list = page.get_images(full=True)
            
            # Process images if any
            for img_index, img_info in enumerate(image_list):
                xref = img_info[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                
                # Create a PIL Image from the bytes
                image = Image.open(io.BytesIO(image_bytes))
                
                # Save image to a temporary file
                img_filename = f"image_p{page_num + 1}_{img_index + 1}.png"
                image_path = os.path.join(os.path.dirname(pdf_path), img_filename)
                image.save(image_path)
                
                # Add image reference to markdown
                md_content.append(f"![Image {img_index + 1} from page {page_num + 1}]({img_filename})\n\n")
        
        return "".join(md_content)
    
    except Exception as e:
        return f"Error converting PDF to Markdown: {str(e)}"


def convert_docx_to_md(docx_path, output_path="bill_generator/bill.md"):
    """
    Convert a DOCX file to Markdown format, preserving element order.
    
    Args:
        docx_path (str): Path to the DOCX file
        output_path (str): Path to save the markdown file
    Returns:
        str: Markdown content
    """
    try:
        doc = docx.Document(docx_path) # This is docx.api.Document, not docx.document.Document
        md_content = []

        for block in doc.element.body:
            if isinstance(block, CT_P):
                para = DocxParagraphObject(block, doc) # Use the aliased Paragraph
                if not para.text.strip():
                    if not any(isinstance(run.element, docx.oxml.drawing.CT_Drawing) for run in para.runs): # Keep empty lines if they don't contain only images
                         md_content.append('\\n') # Preserve paragraph breaks more consistently
                    # Check for images within paragraphs
                    for run in para.runs:
                        if isinstance(run.element, docx.oxml.drawing.CT_Drawing):
                            # This part is complex: python-docx doesn't directly give image bytes from CT_Drawing.
                            # It typically involves accessing inline_shapes or shapes.
                            # For simplicity, we'll note an image was here. A more robust solution would extract it.
                            # This requires iterating through doc.inline_shapes and matching, which is non-trivial here.
                            # For now, let's rely on a more dedicated image extraction pass if needed,
                            # or a library that handles this better for docx to md.
                            # md_content.append(f"![Image in paragraph]\\n\\n") # Placeholder
                            pass # Let's assume images are handled by a different mechanism or not primary for now

                elif para.style.name.startswith('Heading'):
                    heading_level = 0
                    try:
                        heading_level = int(para.style.name[-1])
                    except ValueError: # Handle cases like "Heading" without a number
                        if para.style.name.lower() == "title":
                            heading_level = 1
                        elif para.style.name.lower() == "subtitle":
                            heading_level = 2
                        else: # Default for unnumbered "Heading" styles
                            heading_level = 1 # Or some other default
                    
                    if heading_level > 0:
                         md_content.append('#' * heading_level + ' ' + para.text.strip() + '\\n\\n')
                    else: # Fallback for styles that start with "Heading" but have no parsable level
                         md_content.append('## ' + para.text.strip() + '\\n\\n')


                else:
                    # Handle bold, italic, etc. within the paragraph run by run
                    line_content = []
                    for run in para.runs:
                        text = run.text
                        if run.bold and run.italic:
                            text = f"***{text}***"
                        elif run.bold:
                            text = f"**{text}**"
                        elif run.italic:
                            text = f"*{text}*"
                        # Add more run-level formatting checks if needed (underline, strikethrough, etc.)
                        line_content.append(text)
                    md_content.append("".join(line_content) + '\\n\\n')

            elif isinstance(block, CT_Tbl):
                table = DocxTableObject(block, doc) # Use the aliased Table
                # Start table
                
                if not table.rows: # Skip empty tables
                    continue

                # Determine column widths for Markdown table alignment (optional, but good for looks)
                # For simplicity, we are not doing this here, but it's a possible enhancement.

                # Header row
                header_cells = table.rows[0].cells
                md_content.append('| ' + ' | '.join(cell.text.strip() for cell in header_cells) + ' |\\n')
                
                # Separator row
                md_content.append('| ' + ' | '.join('---' for _ in header_cells) + ' |\\n')
                
                # Data rows
                for row in table.rows[1:]:
                    md_content.append('| ' + ' | '.join(cell.text.strip() for cell in row.cells) + ' |\\n')
                
                md_content.append('\\n')
        
        # Clean up excessive newlines that might result from the above logic
        final_md = "".join(md_content)
        final_md = re.sub(r'(\\n){3,}', '\\n\\n', final_md) # Replace 3+ newlines with 2
        
        # Save the file to the output path
        try:
            with open(output_path, 'w', encoding='utf-8') as file:
                file.write(final_md)
        except Exception as file_error:
            print(f"Warning: Could not save file to {output_path}: {str(file_error)}")
            
        return final_md.strip()

    except Exception as e:
        return f"Error converting DOCX to Markdown: {str(e)}"


def convert_docx_to_pdf(docx_path, output_pdf_path):
    """
    Convert a DOCX file to PDF format.
    
    Args:
        docx_path (str): Path to the DOCX file.
        output_pdf_path (str): Path to save the generated PDF file.
        
    Returns:
        bool: True if conversion was successful, False otherwise.
    """
    try:
        docx_to_pdf_converter(docx_path, output_pdf_path)
        return True
    except Exception as e:
        # docx2pdf can raise various errors, including if Word/LibreOffice is not found
        # or if there are issues with the document itself.
        error_message = str(e)
        if "Neither MS Word nor LibreOffice found" in error_message:
            print("Error converting DOCX to PDF: Microsoft Word or LibreOffice not found. Please install one and ensure it's in your PATH.")
        else:
            print(f"Error converting DOCX to PDF: {error_message}")
        return False


def convert_txt_to_md(txt_path):
    """
    Convert a TXT file to Markdown format.
    
    Args:
        txt_path (str): Path to the TXT file
        
    Returns:
        str: Markdown content
    """
    try:
        with open(txt_path, 'r', encoding='utf-8') as file:
            content = file.read()
        
        # Simple heuristics to identify structure
        lines = content.split('\n')
        md_content = []
        
        for line in lines:
            # Skip empty lines
            if not line.strip():
                md_content.append('\n')
                continue
            
            # Check if line might be a heading (all caps or ends with colon)
            if line.isupper() or line.strip().endswith(':'):
                md_content.append(f"## {line}\n\n")
            # Check if line starts with a number or bullet
            elif re.match(r'^\d+[\.\)]', line.strip()) or line.strip().startswith('•'):
                md_content.append(f"{line}\n")
            else:
                md_content.append(f"{line}\n")
        
        return "".join(md_content)
    
    except Exception as e:
        return f"Error converting TXT to Markdown: {str(e)}"


def convert_file_to_md(file_path):
    """
    Convert a file to Markdown based on its extension.
    
    Args:
        file_path (str): Path to the file
        
    Returns:
        str: Markdown content
    """
    file_path = Path(file_path)
    extension = file_path.suffix.lower()
    
    if extension == '.pdf':
        return convert_pdf_to_md(file_path)
    elif extension == '.docx':
        return convert_docx_to_md(file_path)
    elif extension == '.txt':
        return convert_txt_to_md(file_path)
    else:
        return f"Unsupported file format: {extension}"


def save_markdown_file(markdown_content, output_path):
    """
    Save markdown content to a file.
    
    Args:
        markdown_content (str): Markdown content to save
        output_path (str): Path to save the markdown file
        
    Returns:
        bool: True if successful, False otherwise
    """
    try:
        with open(output_path, 'w', encoding='utf-8') as file:
            file.write(markdown_content)
        return True
    except Exception as e:
        print(f"Error saving markdown file: {str(e)}")
        return False

def convert_md_to_pdf(md_path, output_path):
    """
    Convert a Markdown file to PDF format with proper rendering.
    
    Args:
        md_path (str): Path to the Markdown file
        output_path (str): Path to save the PDF file
        
    Returns:
        bool: True if successful, False otherwise
    """
    try:
        # Read markdown content
        with open(md_path, 'r', encoding='utf-8') as file:
            md_content = file.read()

        # --- Enhanced logic to find and extract the main Markdown block ---
        cleaned_md_content = md_content # Default to original content
        
        # Try to find the start of a potential ```markdown block or ``` block
        block_start_variants = ["```markdown", "```"]
        actual_start_marker = None
        start_index = -1

        for marker in block_start_variants:
            idx = md_content.find(marker)
            if idx != -1:
                # Ensure we don't select a marker that's part of a larger word or inside another block
                # This is a simple check; more complex scenarios might need regex
                if idx == 0 or md_content[idx-1].isspace(): 
                    start_index = idx + len(marker) # Start after the marker
                    actual_start_marker = marker
                    break
        
        if actual_start_marker: # If we found a start marker
            # Find the corresponding end marker ```, searching from after the start marker
            end_index = md_content.rfind("```", start_index)
            
            if end_index != -1:
                # Extract the content between the markers
                extracted_block = md_content[start_index:end_index].strip()
                cleaned_md_content = extracted_block
            else:
                # If no end marker, but we had a start, it could be that the LLM just prepended text
                # and the rest is the document without a final fence. This is a heuristic.
                # For now, if it started with a fence but didn't end with one, we might have an issue.
                # Let's assume for now that if a start marker is found, an end marker should exist for a clean block.
                # If only a start marker is found, the original stripping logic might be safer if it applied.
                # Fallback to simple stripping if the complex extraction isn't clean.
                stripped_md = md_content.strip()
                if stripped_md.startswith(actual_start_marker) and stripped_md.endswith("```"):
                    lines = stripped_md.split('\n')
                    if lines:
                        lines.pop(0)
                    if lines and lines[-1].strip() == "```":
                        lines.pop(-1)
                    cleaned_md_content = "\n".join(lines)
                # else: cleaned_md_content remains md_content or the result of the more specific extraction
        else:
            # No ```markdown or ``` start found, try simple stripping for just ```
            stripped_md = md_content.strip()
            if stripped_md.startswith("```") and stripped_md.endswith("```"):
                lines = stripped_md.split('\n')
                if lines: lines.pop(0) # remove first line
                if lines and lines[-1].strip() == "```": lines.pop(-1) # remove last line
                cleaned_md_content = "\n".join(lines)
        # --- End of enhanced logic ---

        # --- DEBUG PRINT for cleaned_md_content --- #
        print("--- Cleaned MD Content for Markdown Parser ---")
        print(repr(cleaned_md_content)) # Use repr() to see exact string with newlines
        print("---------------------------------------------")

        # Convert cleaned markdown to an HTML snippet
        html_snippet = markdown.markdown(cleaned_md_content, extensions=['tables', 'fenced_code'])

        # Prepend letterhead spacer div to the HTML snippet
        letterhead_spacer_html = "<div style=\"height: 3cm;\"></div>" # Approx 2cm spacer
        html_snippet_with_spacer = letterhead_spacer_html + html_snippet

        # --- DEBUG PRINT for HTML snippet --- #
        print("--- Generated HTML Snippet for PDF Conversion ---")
        print(html_snippet_with_spacer) # Print the one with the spacer
        print("--------------------------------------------------")

        # Define CSS for styling
        styling_css = """
            body { font-family: Arial, sans-serif; margin: 0; line-height: 1.2; } /* Adjusted margin for snippet */
            h1 { color: #333; margin-bottom: 10px; font-size: 24px; }
            h2 { color: #444; margin-bottom: 8px; font-size: 20px; }
            h3 { color: #555; margin-bottom: 6px; font-size: 16px; }
            p { margin-top: 5px; margin-bottom: 5px; }
            pre { background-color: #f5f5f5; padding: 8px; border-radius: 5px; white-space: pre-wrap; word-wrap: break-word; }
            code { font-family: monospace; }
            blockquote { border-left: 4px solid #ccc; padding-left: 15px; color: #777; margin: 10px 0; }
            table { border-collapse: collapse; width: 100%; margin-bottom: 10px; }
            th, td { border: 1px solid #ddd; padding: 6px; text-align: left; }
            th { background-color: #f2f2f2; }
            ul, ol { margin-top: 5px; margin-bottom: 5px; padding-left: 20px; }
            li { margin-bottom: 3px; }
        """
        
        # Create PDF document
        doc = fitz.open()
        
        # Add a new page with A4 dimensions
        page = doc.new_page(width=595, height=842) # A4 dimensions in points
        
        # Define the rectangle for content insertion (with margins)
        # Margins: 50 points top/bottom, 50 points left/right
        rect = fitz.Rect(50, 50, 595 - 50, 842 - 50)
        
        # Use insert_htmlbox with the HTML snippet and the separate CSS string
        # Wrap the snippet in a basic body tag for context if needed by fitz, though often not necessary for snippets
        # However, applying body styles from CSS might need it.
        # Let's try with a minimal body wrapper.
        html_for_fitz = f"""<!DOCTYPE html><html><head><meta charset="utf-8"></head>
        <body>{html_snippet_with_spacer}</body></html>""" # Use the snippet with spacer

        # The 'archive' parameter is for resource management (e.g., external images referenced in HTML)
        # For self-contained HTML like ours, None should be appropriate.
        archive_param = None 

        returned_value = page.insert_htmlbox(rect, html_for_fitz, css=styling_css, archive=archive_param)
        
        # Check for specific error codes or overflow
        if returned_value == -1:
            print("Warning: fitz.insert_htmlbox general error (code -1).")
        elif returned_value == -2:
            print("Warning: fitz.insert_htmlbox invalid HTML/CSS (code -2).")
        elif isinstance(returned_value, fitz.Rect) and not returned_value.is_empty:
            # This means some content did not fit into the provided rectangle (overflow)
            print(f"Warning: fitz.insert_htmlbox indicates content overflow. Overflow area: {returned_value}")
        # If returned_value is a fitz.Rect and is_empty is True, it means all content fitted.
        # No specific message for this case, as it's the expected success.

        # Save PDF
        doc.save(output_path)
        doc.close()
        
        # Redundant temporary HTML file handling is removed.
        
        return True
    except Exception as e:
        print(f"Error converting Markdown to PDF: {str(e)}")
        return False


def convert_md_to_docx(md_path, output_path):
    """
    Convert a Markdown file to DOCX format.
    
    Args:
        md_path (str): Path to the Markdown file
        output_path (str): Path to save the DOCX file
        
    Returns:
        bool: True if successful, False otherwise
    """
    try:
        # Read markdown content
        with open(md_path, 'r', encoding='utf-8') as file:
            md_content = file.read()
        
        # Convert markdown to HTML
        html_content = markdown.markdown(md_content)
        
        # Create a new Word document
        doc = docx.Document()
        
        # Parse HTML and add content to Word document
        # This is a simplified approach - for complex documents, 
        # you might need a more sophisticated HTML parser
        for line in html_content.split('\n'):
            # Remove HTML tags (simplified approach)
            clean_line = re.sub(r'<.*?>', '', line).strip()
            if clean_line:
                doc.add_paragraph(clean_line)
        
        # Save the document
        doc.save(output_path)
        
        return True
    except Exception as e:
        print(f"Error converting Markdown to DOCX: {str(e)}")
        return False


def convert_md_to_txt(md_path, output_path):
    """
    Convert a Markdown file to plain text format.
    
    Args:
        md_path (str): Path to the Markdown file
        output_path (str): Path to save the TXT file
        
    Returns:
        bool: True if successful, False otherwise
    """
    try:
        # Read markdown content
        with open(md_path, 'r', encoding='utf-8') as file:
            md_content = file.read()
        
        # Remove markdown formatting (simplified approach)
        # Remove headers
        txt_content = re.sub(r'#+\s+', '', md_content)
        # Remove bold/italic
        txt_content = re.sub(r'\*\*|\*|__|\b_\b', '', txt_content)
        # Remove links
        txt_content = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', txt_content)
        # Remove images
        txt_content = re.sub(r'!\[([^\]]+)\]\([^)]+\)', '', txt_content)
        # Remove code blocks
        txt_content = re.sub(r'```[^`]*```', '', txt_content)
        
        # Save as text file
        with open(output_path, 'w', encoding='utf-8') as file:
            file.write(txt_content)
        
        return True
    except Exception as e:
        print(f"Error converting Markdown to TXT: {str(e)}")
        return False


def convert_md_to_file(md_path, output_path):
    """
    Convert a Markdown file to another format based on the output extension.
    
    Args:
        md_path (str): Path to the Markdown file
        output_path (str): Path to save the output file
        
    Returns:
        bool: True if successful, False otherwise
    """
    output_path = Path(output_path)
    extension = output_path.suffix.lower()
    
    if extension == '.pdf':
        return convert_md_to_pdf(md_path, output_path)
    elif extension == '.docx':
        return convert_md_to_docx(md_path, output_path)
    elif extension == '.txt':
        return convert_md_to_txt(md_path, output_path)
    else:
        print(f"Unsupported output format: {extension}")
        return False

if __name__ == "__main__":
    convert_docx_to_pdf("bill_generator/ayur.docx", "bill_generator/bill.pdf")